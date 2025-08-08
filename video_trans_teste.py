import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import whisper
from docx import Document
import os
import threading
from datetime import datetime
import traceback
import subprocess
import sys

class VideoTranscriberApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Video Transcriber - Whisper")
        self.root.geometry("550x400")
        
        # Variáveis de controle
        self.file_path = tk.StringVar()
        self.status = tk.StringVar(value="Pronto")
        self.progress = tk.DoubleVar()
        self.model_size = tk.StringVar(value="base")
        self.model = None
        
        # Configura a interface
        self.create_widgets()
    
    def create_widgets(self):
        main_frame = tk.Frame(self.root, padx=20, pady=20)
        main_frame.pack(expand=True, fill=tk.BOTH)
        
        tk.Label(main_frame, text="Video Transcriber", font=("Arial", 16, "bold")).pack(pady=5)
        
        # Seletor de arquivo com suporte a .MOD
        file_frame = tk.Frame(main_frame)
        file_frame.pack(fill=tk.X, pady=10)
        
        tk.Label(file_frame, text="Vídeo:").pack(side=tk.LEFT)
        tk.Entry(file_frame, textvariable=self.file_path, width=40).pack(side=tk.LEFT, padx=5, expand=True, fill=tk.X)
        tk.Button(file_frame, text="Selecionar", command=self.select_file).pack(side=tk.LEFT)
        
        # Opções de qualidade
        quality_frame = tk.LabelFrame(main_frame, text="Qualidade de Transcrição", padx=10, pady=10)
        quality_frame.pack(fill=tk.X, pady=10)
        
        qualities = [
            ("Básica (rápida, menos precisa)", "base"),
            ("Média (equilíbrio)", "small"),
            ("Aprofundada (lenta, mais precisa)", "medium")
        ]
        
        for text, mode in qualities:
            tk.Radiobutton(quality_frame, text=text, variable=self.model_size, 
                         value=mode, anchor="w").pack(fill=tk.X, pady=2)
        
        tk.Button(main_frame, text="Transcrever Vídeo", command=self.start_transcription_thread,
                height=2, bg="#4CAF50", fg="white", font=("Arial", 10, "bold")).pack(fill=tk.X, pady=10)
        
        ttk.Progressbar(main_frame, variable=self.progress, maximum=100).pack(fill=tk.X, pady=5)
        tk.Label(main_frame, textvariable=self.status, fg="blue", font=("Arial", 9)).pack()
        tk.Label(main_frame, text="© 2023 Video Transcriber - Whisper", fg="gray").pack(side=tk.BOTTOM)
    
    def select_file(self):
        # Atualizado para incluir .MOD e outros formatos profissionais
        filetypes = (
            ('Vídeos profissionais', '*.mod *.MOD *.mp4 *.avi *.mov *.mkv *.wmv *.mxf *.mts'),
            ('Arquivos .MOD', '*.mod *.MOD'),
            ('Todos os arquivos', '*.*')
        )
        
        filename = filedialog.askopenfilename(
            title="Selecione um vídeo",
            initialdir=os.path.expanduser("~/Videos"),
            filetypes=filetypes
        )
        
        if filename:
            self.file_path.set(filename)
            self.status.set(f"Pronto para transcrever (Qualidade: {self.model_size.get().upper()})")
    
    def start_transcription_thread(self):
        if not self.file_path.get():
            messagebox.showwarning("Aviso", "Selecione um arquivo de vídeo primeiro!")
            return
        
        self.root.children["!frame"].children["!button"].config(state=tk.DISABLED)
        self.status.set("Iniciando transcrição...")
        
        thread = threading.Thread(target=self.transcribe_video, daemon=True)
        thread.start()
    
    def format_time(self, seconds):
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        seconds = seconds % 60
        return f"{hours:02d}:{minutes:02d}:{seconds:06.3f}".replace('.', ',')
    
    def convert_mod_to_mp4(self, mod_path):
        """Converte arquivos .MOD para .mp4 usando FFmpeg (se necessário)"""
        try:
            if not mod_path.lower().endswith('.mod'):
                return mod_path
                
            self.status.set("Convertendo .MOD para formato compatível...")
            output_path = os.path.splitext(mod_path)[0] + "_converted.mp4"
            
            cmd = [
                'ffmpeg',
                '-i', mod_path,
                '-c:v', 'libx264',
                '-crf', '23',
                '-preset', 'fast',
                '-c:a', 'aac',
                '-b:a', '192k',
                output_path
            ]
            
            subprocess.run(cmd, check=True, capture_output=True)
            return output_path
            
        except subprocess.CalledProcessError as e:
            raise Exception(f"Falha na conversão do .MOD:\n{e.stderr.decode()}")
        except Exception as e:
            raise Exception(f"Erro inesperado na conversão: {str(e)}")
    
    def transcribe_video(self):
        try:
            # Verificação do arquivo
            video_path = self.file_path.get()
            if not video_path:
                raise Exception("Nenhum arquivo selecionado")
                
            if not os.path.exists(video_path):
                raise Exception(f"Arquivo não encontrado: {video_path}")
            
            # Conversão de .MOD se necessário
            original_path = video_path
            if video_path.lower().endswith('.mod'):
                video_path = self.convert_mod_to_mp4(video_path)
                self.progress.set(15)
            
            # Carregamento do modelo
            self.status.set(f"Carregando modelo {self.model_size.get()}...")
            self.model = whisper.load_model(self.model_size.get())
            self.progress.set(20)
            
            # Transcrição
            self.status.set("Transcrevendo vídeo...")
            result = self.model.transcribe(video_path)
            self.progress.set(70)
            
            if not result or "segments" not in result:
                raise Exception("Transcrição retornou resultado vazio")
            
            # Geração do documento
            self.status.set("Gerando documento Word...")
            doc = Document()
            
            # Cabeçalho
            doc.add_heading("Transcrição do Vídeo", level=1)
            doc.add_paragraph(f"Arquivo original: {os.path.basename(original_path)}")
            doc.add_paragraph(f"Qualidade: {self.model_size.get().upper()}")
            doc.add_paragraph(f"Data: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}\n")
            
            # Tabela de segmentos
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "Início"
            hdr[1].text = "Fim"
            hdr[2].text = "Texto"
            
            for segment in result["segments"]:
                row = table.add_row().cells
                row[0].text = self.format_time(segment["start"])
                row[1].text = self.format_time(segment["end"])
                row[2].text = segment["text"]
            
            # Salvamento
            output_dir = os.path.join(os.path.expanduser("~"), "Documentos")
            os.makedirs(output_dir, exist_ok=True)
            
            base_name = os.path.splitext(os.path.basename(original_path))[0][:50]
            output_path = os.path.join(output_dir, f"{base_name}_transcricao.docx")
            
            doc.save(output_path)
            self.progress.set(100)
            
            # Limpeza (remove arquivo convertido temporário)
            if video_path != original_path and os.path.exists(video_path):
                os.remove(video_path)
            
            self.status.set(f"✅ Concluído! Arquivo salvo em:\n{output_path}")
            messagebox.showinfo("Sucesso", f"Transcrição concluída!\n{output_path}")
        
        except Exception as e:
            self.status.set(f"Erro: {str(e)}")
            messagebox.showerror("Erro", f"Ocorreu um erro:\n{str(e)}")
            print(f"\n=== ERRO DETALHADO ===\n{traceback.format_exc()}\n")
        
        finally:
            self.root.after(0, lambda: self.root.children["!frame"].children["!button"].config(state=tk.NORMAL))
            self.progress.set(0)

def check_dependencies():
    try:
        import whisper
        import torch
        return True
    except ImportError:
        print("⚠️ Dependências não encontradas. Instalando automaticamente...")
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", "torch", "whisper", "python-docx"])
            return True
        except:
            return False

if __name__ == "__main__":
    if not check_dependencies():
        input("Pressione Enter para sair após instalação...")
        sys.exit(1)
    
    try:
        subprocess.run(["ffmpeg", "-version"], capture_output=True, check=True)
    except:
        print("❌ FFmpeg não está instalado. Necessário para processar .MOD!")
        if sys.platform == "win32":
            print("Baixe em: https://ffmpeg.org/download.html")
        input("Pressione Enter para sair...")
        sys.exit(1)
    
    root = tk.Tk()
    app = VideoTranscriberApp(root)
    root.mainloop()